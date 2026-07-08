#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
parse_briefing.py - 解析每日赛训简报 docx
  实际格式: Heading段落 + 多个Table
  Table0: 元数据 (对手/日期/赛事)
  Table1: 选图 (我方选图/对方选图)
  Table2+: 战术指令 (每个Table内可能混合多个节标题 + 指令行)
输入: docx 文件路径 (sys.argv[1])
输出: JSON { "match_date": "...", "opponent": "...", "items": [...] }
"""
import sys, json, re, os
from docx import Document


def parse_briefing(filepath):
    doc = Document(filepath)
    match_date = None
    opponent = 'Unknown'
    event_name = None
    our_maps = []
    opp_maps = []
    all_maps = []
    items = []
    sort_order = 0

    # ---- Step 1: 元数据 (Table 0 + 段落) ----
    if doc.tables:
        t0 = doc.tables[0]
        for row in t0.rows:
            cells = [cell.text.strip() for cell in row.cells]
            label = cells[0] if cells else ''
            val = cells[1] if len(cells) > 1 else ''
            if '约战对象' in label or '对手' in label:
                opponent = val or 'Unknown'
            elif '日期' in label or '时间' in label:
                for pat in [r'(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})',
                            r'(\d{4})年(\d{1,2})月(\d{1,2})日']:
                    m = re.search(pat, val)
                    if m:
                        match_date = f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
                        break
            elif '赛事' in label or '比赛' in label:
                event_name = val or None

    # 段落中找日期
    if not match_date:
        for p in doc.paragraphs:
            t = p.text.strip()
            for pat in [r'^(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})$',
                        r'(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})']:
                m = re.search(pat, t)
                if m:
                    match_date = f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
                    break
            if match_date:
                break

    # 文件名备用
    fname = os.path.basename(filepath)
    if not match_date:
        m = re.search(r'(\d{4})', fname)
        if m:
            d = m.group(1)
            try:
                match_date = f"2026-{int(d[0:2]):02d}-{int(d[2:4]):02d}"
            except:
                pass
    if opponent == 'Unknown':
        m = re.search(r'(?:vs[._]?|VS[._]?)([A-Za-z0-9]+)', fname)
        if m:
            opponent = m.group(1).capitalize()

    # ---- Step 2: 选图 (Table 1) ----
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        # Row0 = 表头 ("我方选图", "对方选图"), Row1 = 值
        rows1 = [[cell.text.strip() for cell in row.cells] for row in t1.rows]
        headers1 = rows1[0] if rows1 else []
        values1 = rows1[1] if len(rows1) > 1 else []
        for i, h in enumerate(headers1):
            if '我方' in h:
                if i < len(values1) and values1[i]:
                    our_maps.append(values1[i])
            elif '对方' in h:
                if i < len(values1) and values1[i]:
                    opp_maps.append(values1[i])

    all_maps = our_maps + opp_maps

    # ---- Step 3: 解析指令表格 (Table 2+) 带上下文追踪 ----
    # 上下文跟踪
    cur_map = ''
    cur_side = 'CT'
    cur_round = ''

    SECTION_PATTERN = re.compile(
        r'(?:地图(\d+))?[-\s]*'
        r'(?:T|CT)?[-\s]*'
        r'(长枪|手枪|半起|ECO|强起|强起局|半ECO|半起局|手枪局|长枪局)'
    )

    TEMPLATE_HEADERS = {'纪律指令', '教练补充说明', '注意事项', '通用指令', '额外说明'}

    for ti in range(2, len(doc.tables)):
        table = doc.tables[ti]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

        for row in rows:
            if not row or not any(c for c in row):
                continue

            # 取主要文本（跨列取第一个非空 cell）
            text = ''
            for c in row:
                if c:
                    text = c
                    break
            if not text:
                continue

            # ---- 判断是否为上下文切换标题 ----
            is_section_header = False

            # 模式1: "地图1 - CT手枪局" / "地图2 - T长枪局"
            m = SECTION_PATTERN.search(text)
            if m:
                map_num = m.group(1)
                if map_num:
                    idx = int(map_num) - 1
                    if 0 <= idx < len(all_maps):
                        cur_map = all_maps[idx]
                # 阵营判断（CT优先，避免CT中的T被误匹配）
                if '- CT' in text or 'CT长枪' in text or 'CT手枪' in text:
                    cur_side = 'CT'
                elif '- T' in text or 'T长枪' in text or 'T手枪' in text or 'T半' in text:
                    cur_side = 'T'
                # 局型
                rt = m.group(2)
                rt_map = {'长枪': '长枪', '长枪局': '长枪',
                          '手枪': '手枪', '手枪局': '手枪',
                          '半起': '半起', '半起局': '半起',
                          'ECO': 'ECO', '半ECO': 'ECO',
                          '强起': '强起', '强起局': '强起'}
                cur_round = rt_map.get(rt, rt)
                is_section_header = True

            # 模式2: "强起局战术安排" / "半起局战术安排" (仅在不匹配模式1时尝试)
            if not is_section_header:
                for label, rt in [('强起', '强起'), ('强钢', '强起'), ('半起', '半起'),
                                   ('半ECO', 'ECO'), ('ECO', 'ECO'),
                                   ('长枪', '长枪'), ('手枪', '手枪')]:
                    if label in text and ('战术' in text or '局' in text):
                        cur_map = ''
                        cur_round = rt
                        is_section_header = True
                        break

            # 模式3: 模板标题行 + 纪律/补充说明等重置上下文
            if not is_section_header:
                if text in TEMPLATE_HEADERS or any(kw in text for kw in
                    ['纪律指令', '注意事项', '补充说明', '通用指令']):
                    cur_map = ''
                    cur_round = ''
                    is_section_header = True

            if is_section_header:
                continue

            # ---- 这是实际战术指令 ----
            sort_order += 1
            items.append({
                'map_name': cur_map,
                'team_side': cur_side,
                'tactic_id': '',
                'round_type': cur_round,
                'priority': '一般',
                'instruction': text,
                'notes': '',
                'sort_order': sort_order,
            })

    # ---- 降级: 段落中的战术编号提取 ----
    if not items:
        tac_pattern = re.compile(r'([A-Z0-9]+-[CT]-\w\d{2})')
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            for m in tac_pattern.finditer(text):
                sort_order += 1
                items.append({
                    'map_name': '',
                    'team_side': 'CT' if 'CT' in m.group(1) else 'T',
                    'tactic_id': m.group(1),
                    'round_type': '',
                    'priority': '一般',
                    'instruction': text,
                    'notes': '',
                    'sort_order': sort_order,
                })

    if not items:
        return {'error': '未解析到战术条目，请检查文档格式'}

    return {
        'match_date': match_date,
        'opponent': opponent,
        'event_name': event_name,
        'maps': {'our': our_maps, 'opponent': opp_maps},
        'items': items,
        'items_count': len(items),
    }


if __name__ == '__main__':
    if len(sys.argv) < 2:
        json.dump({'error': '缺少文件路径参数'}, sys.stdout, ensure_ascii=False)
        sys.exit(1)
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
    try:
        result = parse_briefing(sys.argv[1])
        json.dump(result, sys.stdout, ensure_ascii=False)
    except Exception as e:
        json.dump({'error': str(e)}, sys.stdout, ensure_ascii=False)
        sys.exit(1)
